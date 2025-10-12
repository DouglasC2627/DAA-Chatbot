"""
Document processing service for extracting text and metadata from various file formats.

This service handles:
- PDF processing using pypdf
- DOCX processing using python-docx
- TXT/MD plain text processing
- CSV/XLSX spreadsheet processing
- Text extraction with metadata
- Error handling for corrupted files
"""
import csv
import io
import json
import logging
import mimetypes
from pathlib import Path
from typing import Optional, BinaryIO

from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook

from models.document import DocumentType

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessingError(Exception):
    """Raised when document processing fails."""
    pass


class UnsupportedDocumentTypeError(DocumentProcessingError):
    """Raised when document type is not supported."""
    pass


class DocumentProcessorResult:
    """Result object containing extracted text and metadata."""

    def __init__(
        self,
        text: str,
        metadata: dict,
        page_count: Optional[int] = None,
        word_count: Optional[int] = None
    ):
        """
        Initialize document processor result.

        Args:
            text: Extracted text content
            metadata: Document metadata
            page_count: Number of pages (for paginated documents)
            word_count: Approximate word count
        """
        self.text = text
        self.metadata = metadata
        self.page_count = page_count
        self.word_count = word_count or self._calculate_word_count(text)

    @staticmethod
    def _calculate_word_count(text: str) -> int:
        """
        Calculate approximate word count from text.

        Args:
            text: Text content

        Returns:
            Word count
        """
        return len(text.split())


class DocumentProcessor:
    """Service for processing various document formats."""

    # Supported file extensions mapped to DocumentType
    SUPPORTED_EXTENSIONS = {
        '.pdf': DocumentType.PDF,
        '.docx': DocumentType.DOCX,
        '.doc': DocumentType.DOCX,  # Try to process as DOCX
        '.txt': DocumentType.TXT,
        '.md': DocumentType.MD,
        '.csv': DocumentType.CSV,
        '.xlsx': DocumentType.XLSX,
        '.xls': DocumentType.XLSX,  # Try to process as XLSX
    }

    @classmethod
    def get_document_type(cls, filename: str) -> DocumentType:
        """
        Determine document type from filename.

        Args:
            filename: Name of the file

        Returns:
            DocumentType enum value

        Raises:
            UnsupportedDocumentTypeError: If file type is not supported
        """
        extension = Path(filename).suffix.lower()

        if extension not in cls.SUPPORTED_EXTENSIONS:
            raise UnsupportedDocumentTypeError(
                f"Unsupported file type: {extension}. "
                f"Supported types: {', '.join(cls.SUPPORTED_EXTENSIONS.keys())}"
            )

        return cls.SUPPORTED_EXTENSIONS[extension]

    @classmethod
    def get_mime_type(cls, filename: str) -> str:
        """
        Get MIME type for a file.

        Args:
            filename: Name of the file

        Returns:
            MIME type string
        """
        mime_type, _ = mimetypes.guess_type(filename)
        return mime_type or "application/octet-stream"

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """
        Check if file type is supported.

        Args:
            filename: Name of the file

        Returns:
            True if supported, False otherwise
        """
        extension = Path(filename).suffix.lower()
        return extension in cls.SUPPORTED_EXTENSIONS

    def process_document(
        self,
        file_content: BinaryIO,
        filename: str,
        document_type: Optional[DocumentType] = None
    ) -> DocumentProcessorResult:
        """
        Process document and extract text and metadata.

        Args:
            file_content: File content as binary stream
            filename: Original filename
            document_type: Optional document type (auto-detected if not provided)

        Returns:
            DocumentProcessorResult with extracted content

        Raises:
            DocumentProcessingError: If processing fails
            UnsupportedDocumentTypeError: If document type is not supported
        """
        # Determine document type if not provided
        if document_type is None:
            document_type = self.get_document_type(filename)

        logger.info(f"Processing document: {filename} as {document_type.value}")

        try:
            # Route to appropriate processor
            if document_type == DocumentType.PDF:
                return self._process_pdf(file_content, filename)
            elif document_type == DocumentType.DOCX:
                return self._process_docx(file_content, filename)
            elif document_type in (DocumentType.TXT, DocumentType.MD):
                return self._process_text(file_content, filename)
            elif document_type == DocumentType.CSV:
                return self._process_csv(file_content, filename)
            elif document_type == DocumentType.XLSX:
                return self._process_xlsx(file_content, filename)
            else:
                raise UnsupportedDocumentTypeError(
                    f"Processing not implemented for type: {document_type.value}"
                )

        except UnsupportedDocumentTypeError:
            raise
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            raise DocumentProcessingError(
                f"Failed to process document: {str(e)}"
            ) from e

    def _process_pdf(
        self,
        file_content: BinaryIO,
        filename: str
    ) -> DocumentProcessorResult:
        """
        Process PDF document.

        Args:
            file_content: PDF file content
            filename: Original filename

        Returns:
            DocumentProcessorResult with extracted content

        Raises:
            DocumentProcessingError: If PDF processing fails
        """
        try:
            # Read PDF
            pdf_reader = PdfReader(file_content)

            # Extract metadata
            metadata = {
                "filename": filename,
                "type": "pdf",
                "producer": None,
                "creator": None,
                "title": None,
                "author": None,
                "subject": None,
            }

            # Get PDF metadata if available
            if pdf_reader.metadata:
                metadata.update({
                    "producer": pdf_reader.metadata.get('/Producer', None),
                    "creator": pdf_reader.metadata.get('/Creator', None),
                    "title": pdf_reader.metadata.get('/Title', None),
                    "author": pdf_reader.metadata.get('/Author', None),
                    "subject": pdf_reader.metadata.get('/Subject', None),
                })

            # Extract text from all pages
            text_parts = []
            page_count = len(pdf_reader.pages)

            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        # Add page marker for better context
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {str(e)}")
                    continue

            # Combine all text
            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                raise DocumentProcessingError("No text could be extracted from PDF")

            return DocumentProcessorResult(
                text=full_text,
                metadata=metadata,
                page_count=page_count
            )

        except Exception as e:
            raise DocumentProcessingError(f"PDF processing failed: {str(e)}") from e

    def _process_docx(
        self,
        file_content: BinaryIO,
        filename: str
    ) -> DocumentProcessorResult:
        """
        Process DOCX document.

        Args:
            file_content: DOCX file content
            filename: Original filename

        Returns:
            DocumentProcessorResult with extracted content

        Raises:
            DocumentProcessingError: If DOCX processing fails
        """
        try:
            # Load DOCX document
            doc = DocxDocument(file_content)

            # Extract metadata
            metadata = {
                "filename": filename,
                "type": "docx",
                "title": None,
                "author": None,
                "subject": None,
                "created": None,
                "modified": None,
            }

            # Get core properties if available
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata.update({
                    "title": props.title,
                    "author": props.author,
                    "subject": props.subject,
                    "created": str(props.created) if props.created else None,
                    "modified": str(props.modified) if props.modified else None,
                })

            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            # Combine all text
            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                raise DocumentProcessingError("No text could be extracted from DOCX")

            return DocumentProcessorResult(
                text=full_text,
                metadata=metadata,
                page_count=None  # DOCX doesn't have clear page boundaries
            )

        except Exception as e:
            raise DocumentProcessingError(f"DOCX processing failed: {str(e)}") from e

    def _process_text(
        self,
        file_content: BinaryIO,
        filename: str
    ) -> DocumentProcessorResult:
        """
        Process plain text document (TXT, MD).

        Args:
            file_content: Text file content
            filename: Original filename

        Returns:
            DocumentProcessorResult with extracted content

        Raises:
            DocumentProcessingError: If text processing fails
        """
        try:
            # Determine file type
            extension = Path(filename).suffix.lower()
            doc_type = "markdown" if extension == ".md" else "text"

            # Read text with encoding detection
            content_bytes = file_content.read()

            # Try UTF-8 first, then fallback to latin-1
            try:
                text = content_bytes.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text = content_bytes.decode('latin-1')
                except UnicodeDecodeError:
                    # Last resort: ignore errors
                    text = content_bytes.decode('utf-8', errors='ignore')

            # Extract metadata
            metadata = {
                "filename": filename,
                "type": doc_type,
                "encoding": "utf-8",
            }

            if not text.strip():
                raise DocumentProcessingError("Text file is empty")

            return DocumentProcessorResult(
                text=text,
                metadata=metadata,
                page_count=None
            )

        except Exception as e:
            raise DocumentProcessingError(f"Text processing failed: {str(e)}") from e

    def _process_csv(
        self,
        file_content: BinaryIO,
        filename: str
    ) -> DocumentProcessorResult:
        """
        Process CSV document.

        Args:
            file_content: CSV file content
            filename: Original filename

        Returns:
            DocumentProcessorResult with extracted content

        Raises:
            DocumentProcessingError: If CSV processing fails
        """
        try:
            # Read CSV content
            content_bytes = file_content.read()

            # Try UTF-8 first, then fallback to latin-1
            try:
                text_content = content_bytes.decode('utf-8')
            except UnicodeDecodeError:
                text_content = content_bytes.decode('latin-1', errors='ignore')

            # Parse CSV
            csv_reader = csv.reader(io.StringIO(text_content))
            rows = list(csv_reader)

            if not rows:
                raise DocumentProcessingError("CSV file is empty")

            # Get header and data
            header = rows[0] if rows else []
            data_rows = rows[1:] if len(rows) > 1 else []

            # Extract metadata
            metadata = {
                "filename": filename,
                "type": "csv",
                "row_count": len(rows),
                "column_count": len(header),
                "columns": header,
            }

            # Convert to readable text format
            text_parts = []

            # Add header
            if header:
                text_parts.append("Columns: " + ", ".join(header))
                text_parts.append("")

            # Add data rows
            for i, row in enumerate(data_rows, 1):
                if any(cell.strip() for cell in row):  # Skip empty rows
                    row_text = " | ".join(row)
                    text_parts.append(f"Row {i}: {row_text}")

            full_text = "\n".join(text_parts)

            return DocumentProcessorResult(
                text=full_text,
                metadata=metadata,
                page_count=None
            )

        except Exception as e:
            raise DocumentProcessingError(f"CSV processing failed: {str(e)}") from e

    def _process_xlsx(
        self,
        file_content: BinaryIO,
        filename: str
    ) -> DocumentProcessorResult:
        """
        Process XLSX (Excel) document.

        Args:
            file_content: XLSX file content
            filename: Original filename

        Returns:
            DocumentProcessorResult with extracted content

        Raises:
            DocumentProcessingError: If XLSX processing fails
        """
        try:
            # Load workbook
            workbook = load_workbook(file_content, read_only=True, data_only=True)

            # Extract metadata
            metadata = {
                "filename": filename,
                "type": "xlsx",
                "sheet_count": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames,
            }

            # Extract text from all sheets
            text_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]

                # Add sheet header
                text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")

                # Get all rows
                rows = list(sheet.iter_rows(values_only=True))

                if not rows:
                    text_parts.append("(Empty sheet)")
                    continue

                # Get header
                header = rows[0] if rows else []
                data_rows = rows[1:] if len(rows) > 1 else []

                # Add header
                if header:
                    header_text = " | ".join(str(cell) if cell is not None else "" for cell in header)
                    text_parts.append(f"Columns: {header_text}")
                    text_parts.append("")

                # Add data rows (limit to prevent excessive data)
                max_rows = 1000  # Limit rows per sheet
                for i, row in enumerate(data_rows[:max_rows], 1):
                    # Skip completely empty rows
                    if not any(cell is not None and str(cell).strip() for cell in row):
                        continue

                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    text_parts.append(f"Row {i}: {row_text}")

                if len(data_rows) > max_rows:
                    text_parts.append(f"\n(... {len(data_rows) - max_rows} more rows not shown)")

            # Combine all text
            full_text = "\n".join(text_parts)

            if not full_text.strip():
                raise DocumentProcessingError("No data could be extracted from XLSX")

            workbook.close()

            return DocumentProcessorResult(
                text=full_text,
                metadata=metadata,
                page_count=None
            )

        except Exception as e:
            raise DocumentProcessingError(f"XLSX processing failed: {str(e)}") from e

    def validate_file_size(self, file_size: int, max_size: int) -> bool:
        """
        Validate file size against maximum allowed size.

        Args:
            file_size: Size of file in bytes
            max_size: Maximum allowed size in bytes

        Returns:
            True if valid, False otherwise
        """
        return file_size <= max_size

    def extract_text_preview(self, text: str, max_length: int = 500) -> str:
        """
        Extract a preview of the text content.

        Args:
            text: Full text content
            max_length: Maximum length of preview

        Returns:
            Preview string
        """
        if len(text) <= max_length:
            return text

        # Find a good breaking point (end of sentence or word)
        preview = text[:max_length]
        last_period = preview.rfind('.')
        last_space = preview.rfind(' ')

        if last_period > max_length * 0.8:
            return preview[:last_period + 1] + "..."
        elif last_space > max_length * 0.8:
            return preview[:last_space] + "..."
        else:
            return preview + "..."


# Create global document processor instance
document_processor = DocumentProcessor()
